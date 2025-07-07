import streamlit as st
import pandas as pd
import requests
import io
import re
from collections import Counter
import pymorphy3
import base64
from io import BytesIO
from docx import Document
import os
from dotenv import load_dotenv
import urllib3
from requests.adapters import HTTPAdapter
from urllib3.util.retry import Retry
from urllib.parse import urlparse, urljoin
from typing import List, Dict, Tuple, Optional, Any, Set
from bs4 import BeautifulSoup
import pymorphy3
import Stemmer
from rapidfuzz import fuzz
import html
import traceback
import asyncio
import aiohttp

# =================== НАСТРОЙКИ ===================
load_dotenv()
st.set_page_config(page_title="SEO-комбайн", layout="wide", page_icon="🧙‍♂️")

API_TOKEN = "13aef476b8d8f6252959163f5d49812b"
API_URL = f"https://api.serpstat.com/v4?token={API_TOKEN}" if API_TOKEN else None

AUTO_EXTEND = {
    "always": "always прокладки",
    "олвейс": "олвейс прокладки",
    "libresse": "libresse прокладки",
    "nurofen": "nurofen таблетки",
    "но-шпа": "но-шпа таблетки",
    "no-shpa": "no-shpa таблетки",
}

RELEVANT_WORDS = [
    "проклад", "ежеднев", "ночн", "гигиен", "always", "олвейс",
    "таблет", "nurofen", "но-шпа", "libresse", "no-shpa"
]

# --- КОНСТАНТЫ ДЛЯ НОВОЙ ВКЛАДКИ ---
COL_URL_RU_EXCEL = 'URL'
COL_TITLE_RU_EXCEL = 'Title RU'
COL_DESC_RU_EXCEL = 'Description RU'
COL_TITLE_UA_EXCEL = 'Title UA'
COL_DESC_UA_EXCEL = 'Description UA'
COL_EXACT_PHRASES_RU_EXCEL = 'Фразы в точном вхождении RU'
COL_EXACT_PHRASES_UA_EXCEL = 'Фразы в точном вхождении UA'
COL_LSI_RU_EXCEL = 'LSI'
COL_LSI_UA_EXCEL = 'LSI UA'
PREFIX_EXACT_PHRASE = "🔍 "
PREFIX_LSI_PHRASE = "🔤 "
BASE_URL_SITE = 'https://apteka911.ua'
DEFAULT_ENABLE_LSI_TRUNCATION = True
DEFAULT_LSI_TRUNC_MAX_REMOVE = 3
DEFAULT_LSI_TRUNC_MIN_ORIG_LEN = 7
DEFAULT_LSI_TRUNC_MIN_FINAL_LEN = 4
DEFAULT_STEM_FUZZY_RATIO_THRESHOLD = 90

# --- ИНИЦИАЛИЗАЦИЯ ---
urllib3.disable_warnings(urllib3.exceptions.InsecureRequestWarning)
RUSSIAN_STEMMER = Stemmer.Stemmer('russian')
RAPIDFUZZ_AVAILABLE = True
morph = pymorphy3.MorphAnalyzer()


session = requests.Session()
retry_strategy = Retry(total=3, backoff_factor=0.5, status_forcelist=[500, 502, 503, 504])
adapter = HTTPAdapter(max_retries=retry_strategy)
session.mount('http://', adapter)
session.mount('https://', adapter)

# ========== ФУНКЦИИ ИЗ ПРЕДЫДУЩЕГО КОДА ==========
def build_query(name):
    n = name.lower().strip()
    for key in AUTO_EXTEND:
        if key in n:
            return AUTO_EXTEND[key]
    return name

def filter_phrases(phrases, relevant_words=RELEVANT_WORDS, top_n=10):
    filtered = []
    for p in phrases:
        lp = p.lower()
        word_count = len(lp.split())
        if word_count >= 1 and any(word in lp for word in relevant_words):  # Упрощённый фильтр
            filtered.append(p)
        if len(filtered) >= top_n:
            break
    return filtered[:top_n] if filtered else phrases[:top_n]  # Возвращаем хотя бы топ-N, если фильтр пуст

def get_serpstat_phrases_top_filtered(keyword, top_n=10):
    if not API_TOKEN:
        st.error("API-токен Serpstat не настроен. Проверь файл .env.")
        return []
    keyword_query = build_query(keyword)
    payload = {
        "id": "1",
        "method": "SerpstatKeywordProcedure.getKeywords",
        "params": {
            "keyword": keyword_query,
            "se": "g_ua",
            "type": "phrase_all",
            "page": 1,
            "size": 500
        }
    }
    try:
        st.write(f"Запрос к API для ключевого слова: {keyword_query}")
        resp = requests.post(API_URL, json=payload, timeout=30)
        resp.raise_for_status()
        result = resp.json()
        st.write(f"Ответ от API: {result}")
        if "result" in result and "data" in result["result"]:
            data = result["result"]["data"]
            data = [d for d in data if "keyword" in d and "region_queries_count" in d]
            if not data:
                st.write("API вернул пустой список данных.")
                return []
            data.sort(key=lambda d: int(d["region_queries_count"]), reverse=True)
            all_phrases = [d['keyword'] for d in data]
            filtered_phrases = filter_phrases(all_phrases, RELEVANT_WORDS, top_n)
            st.write(f"Отфильтрованные фразы: {filtered_phrases}")
            return filtered_phrases
        else:
            st.write("Нет данных в результате API.")
            return []
    except requests.exceptions.RequestException as e:
        st.error(f"Ошибка подключения к API Serpstat: {str(e)}")
        return []
    except Exception as e:
        st.error(f"Неожиданная ошибка при запросе к API: {str(e)}")
        return []

def analyze_texts(text1: str, text2: str) -> str:
    try:
        morph = pymorphy3.MorphAnalyzer()
    except Exception as e:
        st.error(f"Ошибка при инициализации морфологического анализатора: {e}")
        return "Не удалось выполнить анализ."

    text1_cleaned = re.sub(r'\s+', ' ', text1)
    text2_cleaned = re.sub(r'\s+', ' ', text2)
    def get_lemmas(text: str) -> list:
        words = re.findall(r'\b[а-яА-ЯёЁ-]+\b', text.lower())
        lemmas = [morph.parse(word)[0].normal_form for word in words]
        return lemmas

    lemmas1_set = set(get_lemmas(text1_cleaned))
    all_lemmas_from_text2 = get_lemmas(text2_cleaned)
    unique_lemmas = [lemma for lemma in all_lemmas_from_text2 if lemma not in lemmas1_set]
    if not unique_lemmas:
        return "Во втором тексте не найдено уникальных слов, отсутствующих в первом."
    frequency_counter = Counter(unique_lemmas)
    sorted_lemmas = sorted(frequency_counter.items(), key=lambda x: (-x[1], x[0]))
    result_markdown = [f"- **{word}** — {count} раз(а)" for word, count in sorted_lemmas[:300]]
    return "\n".join(result_markdown)

# ========== НОВАЯ ВКЛАДКА: SEO Meta Checker ==========
def normalize_for_search(text: Optional[Any]) -> str:
    if not text: return ""
    text_str = str(text).lower()
    text_str = re.sub(r'(цена от|ціна від|price from)\s*[^-\|\n\r<]+?(\s*[-\|]|$)', ' ', text_str, flags=re.IGNORECASE)
    text_str = re.sub(r'%[a-z_]+price[a-z_]*%', ' ', text_str)
    text_str = re.sub(r'%[a-z_]+%', ' ', text_str)
    text_str = re.sub(r'\d+(\.\d+)?\s*(грн|uah|usd|eur)?', ' ', text_str)
    text_str = re.sub(r'[/\\|\-–—]', ' ', text_str)
    text_str = re.sub(r'[\s\u00a0\u200b]+', ' ', text_str)
    text_str = re.sub(r'[^\w\s]', '', text_str, flags=re.UNICODE)
    text_str = re.sub(r'\s+', ' ', text_str)
    return text_str.strip()

def split_phrases(phrases_text: Optional[Any]) -> List[str]:
    if not phrases_text or (isinstance(phrases_text, float) and pd.isna(phrases_text)) or not str(phrases_text).strip():
        return []
    phrases = re.split(r'[,\n]+', str(phrases_text))
    return [p.strip() for p in phrases if p.strip()]

@st.cache_data
def cached_lemmatize_word_flexibly(word_to_lemmatize: str, debug_mode_for_messages: bool = False) -> Set[str]:
    clean_word = word_to_lemmatize.strip().lower()
    if not clean_word: return set()
    analyses = morph.analyze(clean_word)
    possible_lemmas = set()
    for analysis_item in analyses:
        if analysis_item.get('analysis') and analysis_item['analysis']:
            lemma = analysis_item['analysis'][0]['lex'].lower()
            if lemma.strip().isalnum(): possible_lemmas.add(lemma)
        elif not possible_lemmas and 'text' in analysis_item:
            original_form = analysis_item['text'].lower()
            if original_form.strip().isalnum(): possible_lemmas.add(original_form)
    if debug_mode_for_messages and not possible_lemmas and clean_word:
        if 'debug_messages' not in st.session_state: st.session_state.debug_messages = []
        debug_msg_detail = f"Анализ Mystem: {str(analyses)[:100]}..." if analyses else "Нет анализа"
        st.session_state.debug_messages.append(f"[LSI Word (Cached)] Слово '{clean_word}' не дало лемм. {debug_msg_detail}")
    return possible_lemmas

def get_primary_lemmas_from_normalized_text(normalized_text: str, debug_mode: bool = False) -> Set[str]:
    if not normalized_text: return set()
    lemmas = morph.lemmatize(normalized_text)
    cleaned_lemmas = {token.lower() for token in lemmas if token.strip().isalnum()}
    if debug_mode:
        if 'debug_messages' not in st.session_state: st.session_state.debug_messages = []
        text_sample = normalized_text[:50] + "..." if len(normalized_text) > 50 else normalized_text
        st.session_state.debug_messages.append(f"[Page Primary Lemmas] Из норм. текста ('{text_sample}', {len(normalized_text.split())} слов) "
                                               f"получено {len(cleaned_lemmas)} уник. осн. лемм. "
                                               f"Пример: {list(cleaned_lemmas)[:10] if cleaned_lemmas else 'Нет'}")
    return cleaned_lemmas

def get_stem_for_word(word: str) -> Optional[str]:
    if not RUSSIAN_STEMMER or not word or not word.strip(): return None
    return RUSSIAN_STEMMER.stemWord(word.lower().strip())

def check_exact_phrases(text: str, phrases_input: Optional[Any], debug_mode: bool = False) -> Dict[str, bool]:
    if not text or not phrases_input or (isinstance(phrases_input, float) and pd.isna(phrases_input)): return {}
    norm_page_text_for_exact_check = normalize_for_search(text)
    phrases_list = split_phrases(phrases_input)
    results = {}
    for phrase in phrases_list:
        norm_phrase_to_find = normalize_for_search(phrase)
        found = False
        if norm_phrase_to_find:
            found = norm_phrase_to_find in norm_page_text_for_exact_check
        if debug_mode:
            if 'debug_messages' not in st.session_state: st.session_state.debug_messages = []
            st.session_state.debug_messages.append({
                'type': 'exact_phrase_debug', 'phrase_original': phrase,
                'details': {'оригинал': phrase, 'нормализованная_фраза': norm_phrase_to_find,
                            'вхождение_в_норм_текст (сэмпл)': norm_page_text_for_exact_check[:100], 'найдено': found}
            })
        results[f"{PREFIX_EXACT_PHRASE}{phrase}"] = found
    return results

def check_lsi_phrases(raw_page_text: str, phrases_input: Optional[Any], debug_mode: bool = False) -> Dict[str, bool]:
    enable_truncation = st.session_state.get('enable_lsi_truncation', DEFAULT_ENABLE_LSI_TRUNCATION)
    trunc_max_remove = st.session_state.get('lsi_trunc_max_remove', DEFAULT_LSI_TRUNC_MAX_REMOVE)
    trunc_min_orig_len = st.session_state.get('lsi_trunc_min_orig_len', DEFAULT_LSI_TRUNC_MIN_ORIG_LEN)
    trunc_min_final_len = st.session_state.get('lsi_trunc_min_final_len', DEFAULT_LSI_TRUNC_MIN_FINAL_LEN)
    current_fuzzy_thresh = st.session_state.get('stem_fuzzy_ratio_threshold', DEFAULT_STEM_FUZZY_RATIO_THRESHOLD)

    if 'debug_messages' not in st.session_state: st.session_state.debug_messages = []

    if not raw_page_text or not phrases_input or (isinstance(phrases_input, float) and pd.isna(phrases_input)):
        if debug_mode: st.session_state.debug_messages.append("### LSI Debug: Нет данных LSI или текст страницы пуст.")
        return {}

    if debug_mode:
        st.session_state.debug_messages.append(f"--- LSI: Входной raw_page_text (первые 300 симв.): ---")
        st.session_state.debug_messages.append(raw_page_text[:300] + "...")
        st.session_state.debug_messages.append(f"(Общая длина raw_page_text: {len(raw_page_text)})")

    normalized_page_content = normalize_for_search(raw_page_text)

    if debug_mode:
        st.session_state.debug_messages.append(f"--- LSI: Текст ПОСЛЕ normalize_for_search (первые 300 симв.): ---")
        st.session_state.debug_messages.append(normalized_page_content[:300] + "...")
        st.session_state.debug_messages.append(f"(Общая длина normalized_page_content: {len(normalized_page_content)}, слов примерно: {len(normalized_page_content.split()) if normalized_page_content else 0})")

    if not normalized_page_content:
        if debug_mode: st.session_state.debug_messages.append("LSI Debug: Текст страницы стал пустым после нормализации. LSI фразы не будут найдены.")
        return {f"{PREFIX_LSI_PHRASE}{p.strip()}": False for p in split_phrases(phrases_input) if p.strip()}

    page_lemmas_set = get_primary_lemmas_from_normalized_text(normalized_page_content, debug_mode)
    page_stems_set = set()
    if RUSSIAN_STEMMER:
        for word_from_norm_page in normalized_page_content.split():
            stemmed_page_word = get_stem_for_word(word_from_norm_page)
            if stemmed_page_word: page_stems_set.add(stemmed_page_word)
        if debug_mode:
            if page_stems_set: st.session_state.debug_messages.append(f"[Page Stems] Найдено {len(page_stems_set)} уник. основ. Пример: {list(page_stems_set)[:10]}")
            elif RUSSIAN_STEMMER: st.session_state.debug_messages.append("[Page Stems] Основы на странице не найдены.")

    phrases_list = split_phrases(phrases_input)
    lsi_results = {}
    for phrase in phrases_list:
        if not phrase.strip(): continue
        normalized_lsi_phrase = normalize_for_search(phrase)
        lsi_phrase_words = [w for w in normalized_lsi_phrase.split() if w]
        if not lsi_phrase_words:
            lsi_results[f"{PREFIX_LSI_PHRASE}{phrase}"] = False
            if debug_mode: st.session_state.debug_messages.append({'LSI_фраза': phrase, 'статус': 'Фраза пуста после нормализации'})
            continue
        all_words_in_lsi_found = True
        debug_lsi_word_checks = []
        for word_in_lsi_phrase in lsi_phrase_words:
            found_this_word = False; match_type = "нет"; matched_page_word_for_display = ""
            possible_lemmas_for_lsi_word = cached_lemmatize_word_flexibly(word_in_lsi_phrase, debug_mode)
            if possible_lemmas_for_lsi_word and any(lemma in page_lemmas_set for lemma in possible_lemmas_for_lsi_word):
                found_this_word = True; match_type = "лемма (гибкая LSI -> основная стр.)"
            lsi_word_stem = None
            if not found_this_word and RUSSIAN_STEMMER:
                lsi_word_stem = get_stem_for_word(word_in_lsi_phrase)
                if lsi_word_stem:
                    if lsi_word_stem in page_stems_set: found_this_word = True; match_type = "основа (точное совпадение)"
                    elif not found_this_word and RAPIDFUZZ_AVAILABLE and page_stems_set:
                        best_fuzzy_match_ratio = 0
                        for page_stem_candidate in page_stems_set:
                            ratio = fuzz.ratio(lsi_word_stem, page_stem_candidate)
                            if ratio >= current_fuzzy_thresh and ratio > best_fuzzy_match_ratio:
                                best_fuzzy_match_ratio = ratio; found_this_word = True
                        if found_this_word: match_type = f"основа (нечеткое, схожесть {best_fuzzy_match_ratio:.0f}%)"
            truncation_info_for_debug = None
            if not found_this_word and enable_truncation:
                original_len = len(word_in_lsi_phrase)
                if original_len >= trunc_min_orig_len:
                    for chars_to_remove in range(1, trunc_max_remove + 1):
                        current_truncated_len = original_len - chars_to_remove
                        if current_truncated_len < trunc_min_final_len: break
                        truncated_lsi_word = word_in_lsi_phrase[:current_truncated_len]
                        try:
                            pattern = r'\b' + re.escape(truncated_lsi_word) + r'\w*\b'
                            match_object = re.search(pattern, normalized_page_content)
                            if match_object:
                                found_this_word = True; matched_page_word_for_display = match_object.group(0)
                                match_type = f"усечение до '{truncated_lsi_word}' (найдено: '{matched_page_word_for_display}')"
                                truncation_info_for_debug = {'truncated_to': truncated_lsi_word, 'matched_on_page': matched_page_word_for_display}; break
                        except re.error as e_re:
                            if debug_mode: st.session_state.debug_messages.append(f"Ошибка Regex (усечение) для '{truncated_lsi_word}': {e_re}")
                            continue
            if debug_mode:
                debug_entry = {'слово_LSI_фразы': word_in_lsi_phrase, 'леммы_LSI_слова': list(possible_lemmas_for_lsi_word),
                               'основа_LSI_слова': lsi_word_stem if lsi_word_stem else (get_stem_for_word(word_in_lsi_phrase) if RUSSIAN_STEMMER else "N/A"),
                               'найдено': found_this_word, 'тип_совпадения': match_type}
                if truncation_info_for_debug: debug_entry['детали_усечения'] = truncation_info_for_debug
                debug_lsi_word_checks.append(debug_entry)
            if not found_this_word: all_words_in_lsi_found = False; break
        lsi_results[f"{PREFIX_LSI_PHRASE}{phrase}"] = all_words_in_lsi_found
        if debug_mode:
            st.session_state.debug_messages.append({'type': 'lsi_phrase_full_debug', 'phrase_original': phrase,
                                                    'details': {'оригинал_LSI_фразы': phrase, 'нормализованная_LSI_фраза': normalized_lsi_phrase,
                                                                'целевые_слова_фразы': lsi_phrase_words, 'детали_проверки_слов': debug_lsi_word_checks,
                                                                'итог_фраза_найдена': all_words_in_lsi_found}})
    return lsi_results

@st.cache_data(ttl=3600)
def get_page_data_for_lang(base_ru_url: str, lang_to_fetch: str, debug_mode_internal: bool = False,
                           save_html_for_debug_manual: bool = False, filename_prefix_manual: str = "manual_debug_page") -> Dict[str, Any]:
    if 'debug_messages' not in st.session_state: st.session_state.debug_messages = []
    headers = {
        'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/125.0.0.0 Safari/537.36',
        'Accept': 'text/html,application/xhtml+xml,application/xml;q=0.9,image/webp,apng,*/*;q=0.8,application/signed-exchange;v=b3;q=0.7',
        'Accept-Language': 'ru-RU,ru;q=1.0,uk;q=0.8,en-US;q=0.6,en;q=0.4' if lang_to_fetch == 'ru' else 'uk-UA,uk;q=1.0,ru;q=0.8,en-US;q=0.6,en;q=0.4',
        'Connection': 'keep-alive', 'Upgrade-Insecure-Requests': '1', 'DNT': '1', 'Sec-GPC': '1',
    }
    session.cookies.set('language', lang_to_fetch, domain='apteka911.ua')
    session.cookies.set('lang', lang_to_fetch, domain='apteka911.ua')

    parsed_original_url = urlparse(base_ru_url)
    original_path = parsed_original_url.path.lstrip('/')
    original_query = parsed_original_url.query

    base_path = original_path
    if base_path.startswith('ua/'):
        base_path = base_path[3:]

    if lang_to_fetch == 'ua':
        final_path_processed = f"ua/{base_path}"
    else:
        final_path_processed = base_path

    base_modified_url = urljoin(BASE_URL_SITE, final_path_processed)
    modified_url_with_query = base_modified_url
    if original_query:
        if not base_modified_url.endswith('?'):
            modified_url_with_query = f"{base_modified_url}?{original_query}"
        else:
            modified_url_with_query = f"{base_modified_url}{original_query}"

    if debug_mode_internal:
        st.session_state.debug_messages.append(f"--- Отладка для URL: {base_ru_url} (язык: {lang_to_fetch}) ---")
        st.session_state.debug_messages.append(f"Исходный URL (base): '{base_ru_url}'")
        st.session_state.debug_messages.append(f"Определен базовый путь: '{base_path}'")
        st.session_state.debug_messages.append(f"Собран финальный URL для запроса: '{modified_url_with_query}'")

    page_title, page_desc, page_full_text, error_message = "", "", "", None
    final_url_after_redirects = modified_url_with_query

    try:
        response = session.get(modified_url_with_query, headers=headers, timeout=20, verify=False, allow_redirects=True)
        response.raise_for_status()
        final_url_after_redirects = response.url
        if debug_mode_internal: st.session_state.debug_messages.append(f"[HTTP RESPONSE] Final URL: '{final_url_after_redirects}', Status: {response.status_code}, Apparent Encoding: {response.apparent_encoding}")

        response.encoding = 'utf-8'
        html_content = response.text

        if save_html_for_debug_manual:
            try:
                path_part = "".join(c if c.isalnum() else "_" for c in urlparse(final_url_after_redirects).path)
                safe_path_part = path_part.replace("__", "_")[:50]
                debug_html_filename = f"{filename_prefix_manual}_{lang_to_fetch}_{safe_path_part}.html"
                with open(debug_html_filename, "w", encoding="utf-8") as f: f.write(html_content)
                if debug_mode_internal: st.session_state.debug_messages.append(f"ОТЛАДКА (ручная): HTML для {final_url_after_redirects} сохранен в: {debug_html_filename}")
            except Exception as e_save:
                if debug_mode_internal: st.session_state.debug_messages.append(f"ОШИБКА РУЧНОГО СОХРАНЕНИЯ HTML: {str(e_save)}")

    except requests.exceptions.RequestException as e:
        error_message = f"Ошибка запроса к {modified_url_with_query}: {str(e)}"
        if debug_mode_internal: st.session_state.debug_messages.append(f"[REQUEST ERROR] URL: {modified_url_with_query}, Error: {error_message}")
        return {'title': '', 'description': '', 'full_text': '', 'error': error_message, 'final_url_fetched': modified_url_with_query}

    soup = BeautifulSoup(html_content, 'html.parser')
    title_tag = soup.find('title')
    description_tag = soup.find('meta', attrs={'name': 'description'})
    page_title = title_tag.text.strip() if title_tag else ''
    page_desc = description_tag['content'].strip() if description_tag and description_tag.get('content') else ''

    if debug_mode_internal:
        st.session_state.debug_messages.append(f"[META EXTRACTED] Из URL: '{final_url_after_redirects}', Title Found: {'Да' if page_title else 'Нет'}, Title: '{page_title[:150]}...'")
        st.session_state.debug_messages.append(f"[META EXTRACTED] Из URL: '{final_url_after_redirects}', Desc Found: {'Да' if page_desc else 'Нет'}, Desc: '{page_desc[:150]}...'")

    content_parts = []
    main_content_selectors = ['article', 'main', '.main-content', '.content', '.post-content', '.entry-content', '#content', '.b-content__body', '.js-mediator-article', '.product-description', '.page-content']
    content_area = None
    for selector in main_content_selectors:
        if selector.startswith('.'): content_area = soup.find(class_=selector[1:])
        elif selector.startswith('#'): content_area = soup.find(id=selector[1:])
        else: content_area = soup.find(selector)
        if content_area:
            if debug_mode_internal and 'debug_messages' in st.session_state: st.session_state.debug_messages.append(f"[CONTENT AREA] Найден по селектору: '{selector}' для URL {final_url_after_redirects}")
            break
    if not content_area and debug_mode_internal and 'debug_messages' in st.session_state: st.session_state.debug_messages.append(f"[CONTENT AREA] Основной блок не найден для {final_url_after_redirects}, используется весь 'soup'.")

    source_tags_container = content_area if content_area else soup
    source_tags = source_tags_container.find_all(['p', 'h1', 'h2', 'h3', 'h4', 'h5', 'h6', 'li', 'span', 'td', 'th', 'strong', 'em', 'b', 'i', 'div'])
    noisy_classes = ['advert', 'social', 'comment', 'sidebar', 'menu', 'nav', 'footer', 'header', 'modal', 'popup', 'banner', 'widget', 'related-posts', 'author-bio', 'breadcrumb', 'pagination', 'meta', 'hidden', 'sr-only', 'price', 'tools', 'actions', 'rating', 'tags', 'share', 'author', 'date', 'category', 'edit-link', 'reply', 'navigation', 'top-link', 'skip-link', 'visually-hidden', 'cookie', 'alert', 'dropdown', 'tab']
    noisy_ids = ['comments', 'sidebar', 'navigation', 'footer', 'header', 'modal', 'popup', 'respond', 'author-info', 'related', 'sharing', 'secondary', 'primary-menu', 'top-bar', 'cookie-banner', 'gdpr-consent']
    noisy_tags_in_parents = ['script', 'style', 'nav', 'footer', 'aside', 'header', 'form', 'button', 'select', 'textarea', 'iframe', 'noscript', 'svg', 'figure', 'figcaption', 'address']
    for tag in source_tags:
        is_noisy = False
        if tag.name == 'div' and not tag.find(['p','li','span','h1','h2','h3','h4','h5','h6'], recursive=False) and len(tag.get_text(strip=True)) < 50: is_noisy = True
        if not is_noisy and tag.has_attr('class') and any(cls in noisy_classes for cls in tag.get('class', [])): is_noisy = True
        if not is_noisy and tag.has_attr('id') and any(id_val in noisy_ids for id_val in tag.get('id', [])): is_noisy = True
        if not is_noisy:
            for parent in tag.parents:
                if parent.name in noisy_tags_in_parents: is_noisy = True; break
                if parent.has_attr('class') and any(cls in noisy_classes for cls in parent.get('class', [])): is_noisy = True; break
                if parent.has_attr('id') and any(id_val in noisy_ids for id_val in parent.get('id', [])): is_noisy = True; break
        if not is_noisy:
            tag_text = tag.get_text(separator=' ', strip=True)
            if tag.name in ['span','div'] and len(tag_text.split()) < 3 and not any(c.isdigit() for c in tag_text):
                if len(tag_text) < 15: continue
            if tag_text: content_parts.append(tag_text)
    content = ' '.join(filter(None, content_parts))
    if debug_mode_internal and 'debug_messages' in st.session_state:
        st.session_state.debug_messages.append(f"[CONTENT SAMPLE for {final_url_after_redirects}] '{content[:300]}...' (Всего символов: {len(content)})")
    page_full_text = f"{page_title} {page_desc} {content}"
    return {'title': page_title, 'description': page_desc, 'full_text': page_full_text, 'error': error_message, 'final_url_fetched': final_url_after_redirects}

@st.cache_data
def load_all_pages_data_for_both_langs(dataframe: pd.DataFrame, _debug_mode_global: bool) -> Dict[str, Dict[str, Dict[str, Any]]]:
    all_data = {}
    total_rows = len(dataframe)
    if 'global_progress_text' not in st.session_state: st.session_state.global_progress_text = ""
    if 'global_progress_value' not in st.session_state: st.session_state.global_progress_value = 0.0
    for i, row in dataframe.iterrows():
        base_ru_url = str(row[COL_URL_RU_EXCEL])
        all_data[base_ru_url] = {}
        for lang_idx, lang_code in enumerate(['ru', 'ua']):
            current_op_total = (i * 2) + (lang_idx + 1)
            total_ops_overall = total_rows * 2
            st.session_state.global_progress_text = f"Загрузка: {current_op_total}/{total_ops_overall} ({base_ru_url} - {lang_code.upper()})"
            st.session_state.global_progress_value = current_op_total / total_ops_overall
            if _debug_mode_global: print(st.session_state.global_progress_text)
            manual_debug_url_val = st.session_state.get("manual_debug_url_val", None)
            manual_debug_lang_val = st.session_state.get("manual_debug_lang_val", None)
            save_html_this_iteration = _debug_mode_global and base_ru_url == manual_debug_url_val and lang_code == manual_debug_lang_val
            all_data[base_ru_url][lang_code] = get_page_data_for_lang(base_ru_url, lang_code,
                                                                      debug_mode_internal=_debug_mode_global,
                                                                      save_html_for_debug_manual=save_html_this_iteration,
                                                                      filename_prefix_manual="auto_save_on_load")
    st.session_state.global_progress_text = "Загрузка данных завершена!"
    st.session_state.global_progress_value = 1.0
    if _debug_mode_global: print(st.session_state.global_progress_text)
    return all_data

def display_debug_messages():
    if 'debug_messages' in st.session_state and st.session_state.debug_messages:
        st.sidebar.markdown("--- Отладочные сообщения ---")
        for msg_item in reversed(st.session_state.debug_messages):
            if isinstance(msg_item, dict) and 'type' in msg_item:
                msg_type = msg_item['type']
                phrase = msg_item.get('phrase_original', msg_item.get('phrase', 'N/A'))
                details = msg_item.get('details', msg_item)
                if msg_type == 'exact_phrase_debug':
                    status_icon = "✅" if details.get('найдено') else "❌"
                    st.sidebar.markdown(f"##### {status_icon} Точная: '{details.get('оригинал', phrase)}'")
                    st.sidebar.json({'Нормализ.': details.get('нормализованная_фраза'), 'Найдено': details.get('найдено')})
                elif msg_type == 'lsi_phrase_full_debug':
                    status_icon = "✅" if details.get('итог_фраза_найдена') else "❌"
                    st.sidebar.markdown(f"##### {status_icon} LSI: '{phrase}'")
                    st.sidebar.json(details)
                else:
                    st.sidebar.write(f"Debug Item (type: {msg_type}):"); st.sidebar.json(msg_item)
            elif isinstance(msg_item, str): st.sidebar.markdown(msg_item)
            else: st.sidebar.write(msg_item)
    if st.sidebar.button("Очистить лог отладки", key="clear_debug_button_sidebar"):
        st.session_state.debug_messages = []
        st.rerun()

def run_checks_for_language(lang_to_check: str, df_excel: pd.DataFrame,
                            all_site_data: Dict[str, Dict[str, Dict[str, Any]]],
                            debug_mode: bool):
    if lang_to_check == 'ua':
        expected_title_col = COL_TITLE_UA_EXCEL; expected_desc_col = COL_DESC_UA_EXCEL
        exact_phrases_col = COL_EXACT_PHRASES_UA_EXCEL; lsi_col = COL_LSI_UA_EXCEL
        if expected_title_col not in df_excel.columns and COL_TITLE_RU_EXCEL in df_excel.columns:
            if debug_mode and 'debug_messages' in st.session_state: st.session_state.debug_messages.append(f"ПРЕДУПРЕЖДЕНИЕ (UA): Кол. Title '{COL_TITLE_UA_EXCEL}' нет, использую '{COL_TITLE_RU_EXCEL}'.")
            expected_title_col = COL_TITLE_RU_EXCEL
        if expected_desc_col not in df_excel.columns and COL_DESC_RU_EXCEL in df_excel.columns:
            if debug_mode and 'debug_messages' in st.session_state: st.session_state.debug_messages.append(f"ПРЕДУПРЕЖДЕНИЕ (UA): Кол. Desc '{COL_DESC_UA_EXCEL}' нет, использую '{COL_DESC_RU_EXCEL}'.")
            expected_desc_col = COL_DESC_RU_EXCEL
    else:
        expected_title_col = COL_TITLE_RU_EXCEL; expected_desc_col = COL_DESC_RU_EXCEL
        exact_phrases_col = COL_EXACT_PHRASES_RU_EXCEL; lsi_col = COL_LSI_RU_EXCEL
    required_cols_for_run = [COL_URL_RU_EXCEL, expected_title_col, expected_desc_col]
    missing_cols_in_df = [col for col in required_cols_for_run if col not in df_excel.columns]
    if missing_cols_in_df:
        st.error(f"Для языка {lang_to_check.upper()}: В Excel отсутствуют ОБЯЗАТЕЛЬНЫЕ колонки для проверки мета-тегов: {', '.join(missing_cols_in_df)}. "
                 f"Ожидались: {', '.join(required_cols_for_run)}")
        return
    sub_tab_meta, sub_tab_phrases = st.tabs([f"📋 Общая проверка Title/Description", f"🔍 Проверка фraz"])
    with sub_tab_meta:
        tab1_errors_summary = {'load_error': 0, 'title_mismatch': 0, 'desc_mismatch': 0}
        tab1_processed_rows_data, urls_with_meta_issues_list_tab1 = [], []
        for index, row_from_df in df_excel.iterrows():
            base_url_from_row = str(row_from_df[COL_URL_RU_EXCEL])
            page_lang_specific_data = all_site_data.get(base_url_from_row, {}).get(lang_to_check, {})
            item_details = {'url': base_url_from_row, 'final_url': page_lang_specific_data.get('final_url_fetched', base_url_from_row), 'has_issue': False, 'issue_details': []}
            if page_lang_specific_data.get('error'):
                tab1_errors_summary['load_error'] += 1; item_details['has_issue'] = True
                item_details['issue_details'].append(f"Ошибка загрузки: {page_lang_specific_data['error']}")
            else:
                expected_title = str(row_from_df.get(expected_title_col, "")).strip()
                expected_desc = str(row_from_df.get(expected_desc_col, "")).strip()
                site_title = page_lang_specific_data.get('title', "").strip()
                site_desc = page_lang_specific_data.get('description', "").strip()
                title_match = normalize_for_search(site_title) == normalize_for_search(expected_title) if expected_title else (not site_title)
                desc_match = normalize_for_search(site_desc) == normalize_for_search(expected_desc) if expected_desc else (not site_desc)
                item_details.update({'expected_title': expected_title, 'expected_desc': expected_desc, 'site_title': site_title, 'site_desc': site_desc, 'title_match': title_match, 'desc_match': desc_match})
                if not title_match: tab1_errors_summary['title_mismatch'] += 1; item_details['has_issue'] = True; item_details['issue_details'].append(f'Title не совпадает')
                if not desc_match: tab1_errors_summary['desc_mismatch'] += 1; item_details['has_issue'] = True; item_details['issue_details'].append(f'Desc не совпадает')
            if item_details['has_issue']: urls_with_meta_issues_list_tab1.append(item_details['final_url'])
            tab1_processed_rows_data.append(item_details)
        st.info(f"Ошибок загрузки: {tab1_errors_summary['load_error']} | Несовп. Title: {tab1_errors_summary['title_mismatch']} | Несовп. Desc: {tab1_errors_summary['desc_mismatch']}")
        show_only_meta_errors_cb = st.checkbox("Показать только URL с ошибками", value=True, key=f"show_err_meta_{lang_to_check}")
        for item_m in tab1_processed_rows_data:
            if show_only_meta_errors_cb and not item_m['has_issue']: continue
            exp_icon = "✅" if not item_m['has_issue'] else "⚠️"
            exp_title = f"{exp_icon} {item_m['final_url']}" + (f" ({', '.join(item_m['issue_details'])})" if item_m['has_issue'] else "")
            with st.expander(exp_title, expanded=item_m['has_issue']):
                if "Ошибка загрузки" in "".join(item_m['issue_details']): st.error(f"Не удалось получить данные. {item_m['issue_details'][0]}")
                else:
                    st.markdown(f"""<div class="info-box"><b>Title:</b> {'✅ Совпадает' if item_m['title_match'] else '❌ Не совпадает'}
                                <div class="result-content"><b>Ожидалось ({expected_title_col}):</b> {html.escape(str(item_m['expected_title']))}</div>
                                <div class="result-content"><b>На сайте:</b> {html.escape(str(item_m['site_title']))}</div></div>
                                <div class="info-box"><b>Description:</b> {'✅ Совпадает' if item_m['desc_match'] else '❌ Не совпадает'}
                                <div class="result-content"><b>Ожидалось ({expected_desc_col}):</b> {html.escape(str(item_m['expected_desc']))}</div>
                                <div class="result-content"><b>На сайте:</b> {html.escape(str(item_m['site_desc']))}</div></div>""", unsafe_allow_html=True)
    with sub_tab_phrases:
        prog_bar_phrases = st.progress(0)
        stat_text_phrases = st.empty()
        exact_col_exists = exact_phrases_col in df_excel.columns
        lsi_col_exists = lsi_col in df_excel.columns
        if not exact_col_exists and debug_mode and 'debug_messages' in st.session_state: st.session_state.debug_messages.append(f"ПРЕДУПРЕЖДЕНИЕ ({lang_to_check.upper()}): Кол. точных фраз '{exact_phrases_col}' не найдена.")
        if not lsi_col_exists and debug_mode and 'debug_messages' in st.session_state: st.session_state.debug_messages.append(f"ПРЕДУПРЕЖДЕНИЕ ({lang_to_check.upper()}): Кол. LSI '{lsi_col}' не найдена.")
        phrase_results_list = []
        for idx, df_row in df_excel.iterrows():
            base_url_from_row = str(df_row[COL_URL_RU_EXCEL])
            stat_text_phrases.info(f"Анализ фраз для URL {idx + 1}/{len(df_excel)}: {base_url_from_row}")
            prog_bar_phrases.progress((idx + 1) / len(df_excel))
            lang_data = all_site_data.get(base_url_from_row, {}).get(lang_to_check, {})

            final_url_for_display = lang_data.get('final_url_fetched', base_url_from_row)

            if lang_data.get('error'):
                phrase_results_list.append({"URL": final_url_for_display, "Тип фразы": "N/A", "Фраза": "Ошибка загрузки страницы", "Статус": "❌ Ошибка"})
                if debug_mode and 'debug_messages' in st.session_state: st.session_state.debug_messages.append(f"Tab2 ({lang_to_check.upper()}): Пропуск фраз {base_url_from_row}, ошибка: {lang_data.get('error','') if lang_data else ''}")
                continue

            full_text = lang_data.get('full_text', "")
            if not full_text and debug_mode and 'debug_messages' in st.session_state: st.session_state.debug_messages.append(f"Tab2 ({lang_to_check.upper()}): Пустой текст для {final_url_for_display}")

            current_url_phrases = {}
            if exact_col_exists and pd.notna(df_row.get(exact_phrases_col)):
                current_url_phrases.update(check_exact_phrases(full_text, df_row[exact_phrases_col], debug_mode))
            if lsi_col_exists and pd.notna(df_row.get(lsi_col)):
                if debug_mode and 'debug_messages' in st.session_state:
                    st.session_state.debug_messages.append(f"--- LSI для URL: {final_url_for_display} ({lang_to_check.upper()}) (Fuzzy: {st.session_state.get('stem_fuzzy_ratio_threshold', DEFAULT_STEM_FUZZY_RATIO_THRESHOLD)}%) ---")
                    st.session_state.debug_messages.append(f"LSI из '{lsi_col}': '{df_row.get(lsi_col)}'")
                current_url_phrases.update(check_lsi_phrases(full_text, df_row[lsi_col], debug_mode))

            for phrase_key, found in current_url_phrases.items():
                p_type = "Точное вхождение" if phrase_key.startswith(PREFIX_EXACT_PHRASE) else "LSI фраза"
                act_phr = phrase_key.replace(PREFIX_EXACT_PHRASE, "").replace(PREFIX_LSI_PHRASE, "")
                phrase_results_list.append({"URL": final_url_for_display, "Тип фразы": p_type, "Фраза": act_phr, "Статус": "✅ Найдено" if found else "❌ Не найдено"})

        stat_text_phrases.success(f"Анализ фраз завершен!")
        prog_bar_phrases.empty()

        if phrase_results_list:
            phrases_df = pd.DataFrame(phrase_results_list)
            urls_w_failed_phr = phrases_df[phrases_df['Статус'].str.contains("❌")]['URL'].nunique()
            total_phr_not_found = len(phrases_df[phrases_df['Статус'] == "❌ Не найдено"])
            st.info(f"URL с ненайденными/ошибочными фразами: {urls_w_failed_phr} | Всего фраз не найдено: {total_phr_not_found}")

            fcols = st.columns(3)
            unique_urls = sorted(phrases_df["URL"].unique().tolist()) if not phrases_df.empty else []
            unique_types = sorted(phrases_df["Тип фразы"].unique().tolist()) if not phrases_df.empty else []
            unique_statuses = sorted(phrases_df["Статус"].unique().tolist()) if not phrases_df.empty else []
            with fcols[0]: url_f = st.multiselect("URL", unique_urls, key=f"url_f_t2_multi_{lang_to_check}")
            with fcols[1]: ptype_f = st.multiselect("Тип фразы", unique_types, key=f"ptype_f_t2_multi_{lang_to_check}")
            with fcols[2]: stat_f = st.selectbox("Статус", ["Все статусы"] + unique_statuses, key=f"stat_f_t2_select_{lang_to_check}")

            filtered_df = phrases_df.copy()
            if url_f: filtered_df = filtered_df[filtered_df["URL"].isin(url_f)]
            if ptype_f: filtered_df = filtered_df[filtered_df["Тип фразы"].isin(ptype_f)]
            if stat_f != "Все статусы": filtered_df = filtered_df[filtered_df["Статус"] == stat_f]

            st.dataframe(filtered_df, hide_index=True, use_container_width=True, height=600,
                         column_config={"URL": st.column_config.TextColumn("Проверенный URL", width="medium"),
                                        "Фраза": st.column_config.TextColumn("Фраза", width="large")})
            csv_dl = filtered_df.to_csv(index=False, encoding='utf-8-sig')
            st.download_button(f"📥 Скачать результаты ({lang_to_check.upper()})", csv_dl, f'filtered_phrases_{lang_to_check}.csv', 'text/csv', key=f"dl_phr_t2_btn_{lang_to_check}")
        else: st.warning(f"Нет данных по фразам для отображения.")

# --- КОНСТАНТЫ ДЛЯ ПРОВЕРКИ ИЗОБРАЖЕНИЙ АПТЕК ---
ID_COLUMN_NAME_EXCEL = "id"
URL_TEMPLATES = [
    "https://tmcu.lll.org.ua/pharmacy_properties_api/files/pharmacies/{ID}/imageApteka.jpeg",
    "https://tmcu.lll.org.ua/pharmacy_properties_api/files/pharmacies/{ID}/imageApteka.png"
]
DEFAULT_USER_AGENT = 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/125.0.0.0 Safari/537.36'

async def fetch_url_status(
    http_session: aiohttp.ClientSession, 
    pharmacy_id: any, 
    url_to_check: str, 
    original_extension_type: str, 
    semaphore: asyncio.Semaphore
) -> dict:
    async with semaphore:
        generated_url = url_to_check
        final_url = ""
        status_code = None
        status_message = "Неизвестная ошибка"
        error_details = ""
        headers = {'User-Agent': DEFAULT_USER_AGENT}
        try:
            async with http_session.get(url_to_check, headers=headers, timeout=aiohttp.ClientTimeout(total=20), ssl=False) as response:
                final_url = str(response.url)
                status_code = response.status
                if 200 <= status_code < 300: status_message = "ОК" if generated_url == final_url else "Редирект ОК"
                elif status_code == 404: status_message = "Не найдено (404)"
                elif status_code == 403: status_message = "Запрещено (403)"
                elif 400 <= status_code < 500: status_message = f"Ошибка клиента ({status_code})"
                elif 500 <= status_code < 600: status_message = f"Ошибка сервера ({status_code})"
                else: status_message = f"Другой статус ({status_code})"
        except asyncio.TimeoutError: status_message = "Таймаут"; error_details = "Запрос > 20 сек"; final_url = generated_url
        except aiohttp.ClientConnectorError as e: status_message = "Ошибка соединения"; error_details = str(e); final_url = generated_url
        except aiohttp.ClientError as e: status_message = "Ошибка клиента (aiohttp)"; error_details = str(e); final_url = generated_url
        except Exception as e: status_message = "Непредвиденная ошибка"; error_details = str(e); final_url = generated_url
        await asyncio.sleep(0.05)
        return {
            "ID аптеки (исходный)": pharmacy_id,
            "Тип файла (проверенный)": original_extension_type.upper(),
            "Проверяемый URL": generated_url,
            "Конечный URL проверки": final_url,
            "HTTP Статус проверки": status_code,
            "Результат проверки URL": status_message,
            "Детали ошибки URL": error_details
        }

async def run_all_checks_async(
    pharmacy_ids_with_raw_values: list,
    url_templates_list: list,
    max_concurrent: int,
    progress_bar_ui,
    status_text_ui
) -> list:
    all_individual_check_results = []
    semaphore = asyncio.Semaphore(max_concurrent)
    connector = aiohttp.TCPConnector(ssl=False)
    async with aiohttp.ClientSession(connector=connector) as http_session:
        tasks = []
        for pharmacy_id_raw, cleaned_id_str_or_none in pharmacy_ids_with_raw_values:
            if cleaned_id_str_or_none is None:
                actual_id_for_report = str(pharmacy_id_raw) if not pd.isna(pharmacy_id_raw) else "ПУСТОЙ_ИЛИ_NAN_ID"
                for template in url_templates_list:
                    ext_type = "JPEG" if ".jpeg" in template.lower() else "PNG" if ".png" in template.lower() else "UNKNOWN"
                    all_individual_check_results.append({
                        "ID аптеки (исходный)": actual_id_for_report,
                        "Тип файла (проверенный)": ext_type,
                        "Проверяемый URL": "N/A (пустой или невалидный ID)",
                        "Конечный URL проверки": "N/A",
                        "HTTP Статус проверки": None,
                        "Результат проверки URL": "Пропущено (ID невалиден)",
                        "Детали ошибки URL": ""
                    })
                continue
            for template in url_templates_list:
                url_to_check = template.replace("{ID}", cleaned_id_str_or_none)
                extension_type = "JPEG" if ".jpeg" in template.lower() else "PNG" if ".png" in template.lower() else "UNKNOWN"
                tasks.append(fetch_url_status(http_session, pharmacy_id_raw, url_to_check, extension_type, semaphore))
        total_tasks_to_run = len(tasks)
        processed_tasks_count = 0
        if progress_bar_ui: progress_bar_ui.progress(0.0)
        if status_text_ui: status_text_ui.text(f"Обработано URL: 0/{total_tasks_to_run} (Всего ID для проверки: {len(pharmacy_ids_with_raw_values)})")
        for i, future in enumerate(asyncio.as_completed(tasks)):
            result_item = None
            try:
                result_item = await future
                all_individual_check_results.append(result_item)
            except Exception as e_task:
                all_individual_check_results.append({
                    "ID аптеки (исходный)": f"Ошибка асинхр. задачи (неизвестный ID)",
                    "Тип файла (проверенный)": "ERROR",
                    "Проверяемый URL": "N/A", "Конечный URL проверки": "N/A",
                    "HTTP Статус проверки": None, "Результат проверки URL": "Ошибка выполнения задачи",
                    "Детали ошибки URL": str(e_task)
                })
            processed_tasks_count += 1
            if progress_bar_ui: progress_bar_ui.progress(processed_tasks_count / total_tasks_to_run if total_tasks_to_run > 0 else 0)
            if status_text_ui:
                id_disp = result_item.get('ID аптеки (исходный)', f'задача {i+1}') if result_item else f'задача {i+1} (ошибка)'
                status_text_ui.text(f"Проверка URL: {processed_tasks_count}/{total_tasks_to_run} (ID: {id_disp})")
    final_aggregated_results = []
    results_by_pharmacy_id = {}
    for res_item in all_individual_check_results:
        pid_raw = res_item.get("ID аптеки (исходный)")
        pid_group_key = str(pid_raw) if not (isinstance(pid_raw, str) and "Ошибка асинхр. задачи" in pid_raw) else f"ERROR_TASK_{time.time()}"
        if pid_group_key not in results_by_pharmacy_id:
            results_by_pharmacy_id[pid_group_key] = {'raw_id': pid_raw, 'checks': []}
        results_by_pharmacy_id[pid_group_key]['checks'].append(res_item)
    for pharmacy_id_key, data in results_by_pharmacy_id.items():
        pharmacy_id_original = data['raw_id']
        checks_for_id = data['checks']
        if len(checks_for_id) == 1 and "Пропущено" in checks_for_id[0].get("Результат проверки URL", ""):
            final_aggregated_results.append({
                "ID аптеки": pharmacy_id_original,
                "Сгенерированный URL (основной)": checks_for_id[0].get("Проверяемый URL", "N/A"),
                "Результат": checks_for_id[0].get("Результат проверки URL", "Ошибка данных"),
                "Найденный формат": checks_for_id[0].get("Тип файла (проверенный)", "N/A"),
                "Конечный URL (найденного)": checks_for_id[0].get("Конечный URL проверки", "N/A"),
                "HTTP Статус (конечный)": checks_for_id[0].get("HTTP Статус проверки", "N/A"),
                "Детали по ошибкам (если есть)": checks_for_id[0].get("Детали ошибки URL", "")
            })
            continue
        jpeg_check = next((c for c in checks_for_id if c.get("Тип файла (проверенный)") == "JPEG"), None)
        png_check = next((c for c in checks_for_id if c.get("Тип файла (проверенный)") == "PNG"), None)
        is_jpeg_ok = jpeg_check and "ОК" in jpeg_check.get("Результат проверки URL", "")
        is_png_ok = png_check and "ОК" in png_check.get("Результат проверки URL", "")
        aggregated_status = "❌ Не найдено (в обоих форматах)"
        found_format_details = "Нет"
        final_url_display = "N/A"
        http_status_display = "N/A"
        error_details_parts = []
        base_id_str_for_template = str(pharmacy_id_original)
        if base_id_str_for_template.endswith(".0"): base_id_str_for_template = base_id_str_for_template[:-2]
        generated_url_template_ref = "N/A (ошибка ID)"
        if base_id_str_for_template and not pd.isna(pharmacy_id_original) and "{ID}" in url_templates_list[0] :
            generated_url_template_ref = url_templates_list[0].replace("{ID}", base_id_str_for_template)
        if is_jpeg_ok and is_png_ok:
            aggregated_status = "✅ ОК"; found_format_details = "JPEG и PNG"
            final_url_display = f"JPEG: {jpeg_check['Конечный URL проверки']}"
            http_status_display = jpeg_check['HTTP Статус проверки']
        elif is_jpeg_ok:
            aggregated_status = "✅ ОК"; found_format_details = "JPEG"
            final_url_display = jpeg_check['Конечный URL проверки']
            http_status_display = jpeg_check['HTTP Статус проверки']
        elif is_png_ok:
            aggregated_status = "✅ ОК"; found_format_details = "PNG"
            final_url_display = png_check['Конечный URL проверки']
            http_status_display = png_check['HTTP Статус проверки']
        else:
            if jpeg_check: error_details_parts.append(f"JPEG: {jpeg_check.get('Результат проверки URL','N/A')} ({jpeg_check.get('Проверяемый URL') or generated_url_template_ref})")
            else: error_details_parts.append(f"JPEG: Проверка не проводилась или ошибка ({generated_url_template_ref})")
            png_template_url_ref = "N/A (ошибка ID)"
            if base_id_str_for_template and not pd.isna(pharmacy_id_original) and len(url_templates_list) > 1 and "{ID}" in url_templates_list[1]:
                png_template_url_ref = url_templates_list[1].replace("{ID}", base_id_str_for_template)
            if png_check: error_details_parts.append(f"PNG: {png_check.get('Результат проверки URL','N/A')} ({png_check.get('Проверяемый URL') or png_template_url_ref})")
            else: error_details_parts.append(f"PNG: Проверка не проводилась или ошибка ({png_template_url_ref})")
            if final_url_display == "N/A": final_url_display = jpeg_check.get('Конечный URL проверки', generated_url_template_ref) if jpeg_check else generated_url_template_ref
            if http_status_display == "N/A": http_status_display = jpeg_check.get('HTTP Статус проверки', "N/A") if jpeg_check else "N/A"
        final_aggregated_results.append({
            "ID аптеки": pharmacy_id_original,
            "Результат": aggregated_status,
            "Найденный формат": found_format_details,
            "Сгенерированный URL (JPEG вариант)": generated_url_template_ref,
            "Конечный URL (если найден)": final_url_display,
            "HTTP Статус (конечный)": http_status_display,
            "Детали по вариантам/ошибкам": "; ".join(error_details_parts) if error_details_parts else ""
        })
    return final_aggregated_results

def pharmacy_image_url_checker_tab():
    st.title("⚕️ Проверка доступности изображений аптек (JPEG и PNG)")
    st.markdown(f"Утилита для массовой проверки URL изображений по ID аптек. Проверяются два варианта: с расширением `.jpeg` и `.png`.")
    uploaded_file = st.file_uploader(
        f"1. Загрузите Excel-файл. Файл должен содержать один столбец с именем **'{ID_COLUMN_NAME_EXCEL}'**, содержащий ID аптек.",
        type=["xlsx", "xls"]
    )
    if uploaded_file:
        try:
            df = pd.read_excel(uploaded_file, dtype={ID_COLUMN_NAME_EXCEL: str})
            st.success(f"Файл '{uploaded_file.name}' успешно загружен. Обнаружено строк: {len(df)}")
            st.markdown("---")
            st.subheader("Предпросмотр данных (первые 5 строк):")
            st.dataframe(df.head())
            if ID_COLUMN_NAME_EXCEL not in df.columns:
                st.error(f"Ошибка: В загруженном файле отсутствует обязательный столбец с именем '{ID_COLUMN_NAME_EXCEL}'.")
            else:
                st.markdown("---")
                st.subheader("2. Настройки проверки")
                max_concurrent_requests = st.slider(
                    "Количество параллельных запросов:",
                    min_value=1, max_value=30, value=10,
                    help="Определяет, сколько URL будет проверяться одновременно."
                )
                if st.button("🚀 Начать проверку!", key="start_check_button"):
                    ids_to_check_raw = df[ID_COLUMN_NAME_EXCEL].tolist()
                    pharmacy_ids_with_raw_values_for_run = []
                    for id_val in ids_to_check_raw:
                        cleaned_id = None
                        if not pd.isna(id_val):
                            id_str = str(id_val).strip()
                            if id_str.endswith(".0"): id_str = id_str[:-2]
                            if id_str: cleaned_id = id_str
                        pharmacy_ids_with_raw_values_for_run.append((id_val, cleaned_id))
                    valid_ids_for_run_count = sum(1 for _, cleaned_id in pharmacy_ids_with_raw_values_for_run if cleaned_id is not None)
                    if not valid_ids_for_run_count:
                        st.warning("Не найдено валидных ID для проверки в файле.")
                    else:
                        st.info(f"Начинается проверка для {len(pharmacy_ids_with_raw_values_for_run)} записей ({valid_ids_for_run_count} ID будут проверены по {len(URL_TEMPLATES)} шаблонам)...")
                        progress_bar_ui = st.progress(0.0)
                        status_text_ui = st.empty()
                        status_text_ui.text("Инициализация...")
                        all_results_aggregated = []
                        try:
                            all_results_aggregated = asyncio.run(
                                run_all_checks_async(pharmacy_ids_with_raw_values_for_run, URL_TEMPLATES, max_concurrent_requests, progress_bar_ui, status_text_ui)
                            )
                        except Exception as e_async_run:
                             st.error(f"Ошибка при выполнении асинхронных задач: {e_async_run}")
                             st.exception(e_async_run)
                        status_text_ui.success(f"Проверка завершена! Обработано записей: {len(all_results_aggregated)}.")
                        if all_results_aggregated:
                            results_df = pd.DataFrame(all_results_aggregated)
                            cols_order = ["ID аптеки", "Результат", "Найденный формат", "Конечный URL (если найден)", "HTTP Статус (конечный)", "Сгенерированный URL (JPEG вариант)", "Детали по вариантам/ошибкам"]
                            final_cols = [col for col in cols_order if col in results_df.columns]
                            results_df_display = results_df[final_cols]
                            st.markdown("---"); st.subheader("📊 Результаты проверки")
                            st.dataframe(results_df_display)
                            output_excel = BytesIO()
                            with pd.ExcelWriter(output_excel, engine='xlsxwriter') as writer:
                                results_df_display.to_excel(writer, index=False, sheet_name='Результаты_проверки_URL')
                                worksheet = writer.sheets['Результаты_проверки_URL']
                                for idx, col in enumerate(results_df_display):
                                    series = results_df_display[col]
                                    max_len = max((series.astype(str).map(len).max(), len(str(series.name)) )) + 2
                                    worksheet.set_column(idx, idx, max_len)
                            excel_data_to_download = output_excel.getvalue()
                            st.download_button(label="📥 Скачать отчет в Excel", data=excel_data_to_download,
                                file_name=f"отчет_URL_изображений_{pd.Timestamp.now().strftime('%Y%m%d_%H%M%S')}.xlsx",
                                mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet")
                        else: st.warning("Результатов для отображения нет.")
        except ValueError as ve:
            st.error(f"Ошибка чтения данных из Excel: {ve}. Убедитесь, что файл корректен и столбец '{ID_COLUMN_NAME_EXCEL}' содержит ожидаемые данные.")
        except Exception as e:
            st.error(f"Произошла ошибка при обработке файла: {e}")
            st.exception(e)

# ========== ОСНОВНАЯ ЛОГИКА ==========
def main():
    # --- Add company logo to sidebar ---
    import os
    logo_path = os.path.join(os.path.dirname(__file__), "WU.png")
    if os.path.exists(logo_path):
        st.sidebar.image(logo_path, width=120)
    st.sidebar.title("🌟 Инструменты SEO-комбайна")
    tab = st.sidebar.radio(
        "Выбери инструмент:",
        [
            "🔎 Подбор ключей (Serpstat)",
            "🧙 Анализ уникальных слов",
            "🔍 Подсветка слов + DOCX",
            "🔍 SEO Meta Checker",
            "🧪 Tittle_Description +",
            "🖼️ Проверка URL изображений аптек",
            "🤖 GPT-ассистент",
            "📖 Инструкция"
        ],
        index=0
    )
    # ========== ВКЛАДКА: GPT-ассистент ==========
    if tab == "🤖 GPT-ассистент":
        import openai
        import time
        from streamlit.components.v1 import html as st_html
        st.title("🤖 GPT-ассистент (OpenAI)")
        st.markdown("""
        <style>
        .gpt-chat-container {
            max-width: 700px;
            margin: 0 auto;
            background: #f7f7f9;
            border-radius: 12px;
            padding: 24px 18px 80px 18px;
            min-height: 400px;
            box-shadow: 0 2px 8px rgba(0,0,0,0.04);
            position: relative;
        }
        .gpt-message {
            display: flex;
            margin-bottom: 18px;
        }
        .gpt-message.user .gpt-bubble {
            background: #e6f0ff;
            color: #222;
            align-self: flex-end;
            margin-left: auto;
        }
        .gpt-message.assistant .gpt-bubble {
            background: #fff;
            color: #222;
            border: 1px solid #e0e0e0;
            align-self: flex-start;
            margin-right: auto;
        }
        .gpt-bubble {
            padding: 12px 16px;
            border-radius: 12px;
            max-width: 80%;
            font-size: 1.08em;
            line-height: 1.6;
            box-shadow: 0 1px 2px rgba(0,0,0,0.03);
            white-space: pre-wrap;
        }
        .gpt-chat-input-bar {
            position: fixed;
            left: 0; right: 0; bottom: 0;
            background: #fff;
            border-top: 1px solid #e0e0e0;
            padding: 18px 0 18px 0;
            z-index: 100;
        }
        .gpt-chat-input-inner {
            max-width: 700px;
            margin: 0 auto;
            display: flex;
            gap: 8px;
        }
        .gpt-chat-input-inner textarea {
            flex: 1;
            border-radius: 8px;
            border: 1px solid #ccc;
            padding: 10px;
            font-size: 1.08em;
            resize: none;
            min-height: 38px;
            max-height: 120px;
        }
        .gpt-chat-input-inner button {
            border-radius: 8px;
            border: none;
            background: #007bff;
            color: #fff;
            font-size: 1.08em;
            padding: 0 18px;
            cursor: pointer;
            transition: background 0.2s;
            height: 38px;
        }
        .gpt-chat-input-inner button:disabled {
            background: #b3d1ff;
            cursor: not-allowed;
        }
        .gpt-chat-actions {
            display: flex;
            gap: 10px;
            margin-bottom: 10px;
            justify-content: flex-end;
        }
        .gpt-chat-actions button {
            background: #f2f2f2;
            color: #333;
            border: 1px solid #ddd;
            border-radius: 6px;
            padding: 4px 12px;
            font-size: 0.98em;
            cursor: pointer;
            transition: background 0.2s;
        }
        .gpt-chat-actions button:hover {
            background: #e6e6e6;
        }
        </style>
        """, unsafe_allow_html=True)

        # --- API KEY (set here, not in UI) ---
        OPENAI_API_KEY = "sk-..."  # <-- ВСТАВЬТЕ СВОЙ КЛЮЧ СЮДА
        openai.api_key = OPENAI_API_KEY

        # --- Session state for chat ---
        if 'gpt_chat_history' not in st.session_state:
            st.session_state.gpt_chat_history = []  # list of {role, content}
        if 'gpt_last_error' not in st.session_state:
            st.session_state.gpt_last_error = None
        if 'gpt_user_input' not in st.session_state:
            st.session_state.gpt_user_input = ""

        # --- Chat actions ---
        col_actions = st.columns([1,1,6])
        with col_actions[0]:
            if st.button("🗑️ Очистить чат", key="gpt_clear_chat_btn"):
                st.session_state.gpt_chat_history = []
                st.session_state.gpt_last_error = None
        with col_actions[1]:
            if st.session_state.gpt_chat_history:
                if st.button("📋 Копировать ответ", key="gpt_copy_btn"):
                    import pyperclip
                    for msg in reversed(st.session_state.gpt_chat_history):
                        if msg['role'] == 'assistant':
                            pyperclip.copy(msg['content'])
                            st.success("Ответ скопирован!")
                            break

        # --- Chat history (top) ---
        st_html('<div class="gpt-chat-container">', height=0)
        for msg in st.session_state.gpt_chat_history:
            role = msg['role']
            bubble_class = 'user' if role == 'user' else 'assistant'
            icon = '🧑' if role == 'user' else '🤖'
            st_html(f'''<div class="gpt-message {bubble_class}"><div class="gpt-bubble">{icon} {msg['content'].replace(chr(10),'<br>')}</div></div>''', height=0)
        if not st.session_state.gpt_chat_history:
            st_html('<div style="color:#888;text-align:center;margin-top:60px;">Нет сообщений. Задайте вопрос!</div>', height=0)
        st_html('</div>', height=0)

        # --- Error display ---
        if st.session_state.gpt_last_error:
            st.error(st.session_state.gpt_last_error)

        # --- Input bar (bottom, fixed) ---
        st_html('''<div class="gpt-chat-input-bar"><div class="gpt-chat-input-inner">''', height=0)
        user_input = st.text_area("", value=st.session_state.gpt_user_input, key="gpt_user_input_area", label_visibility="collapsed", height=70, max_chars=2000, placeholder="Введите ваш вопрос...")
        send_disabled = not user_input.strip() or not OPENAI_API_KEY or len(user_input.strip()) < 1
        send_btn = st.button("Отправить", key="gpt_send_btn", disabled=send_disabled, use_container_width=False)
        st_html('</div></div>', height=0)

        # --- Handle send ---
        if send_btn and not send_disabled:
            st.session_state.gpt_user_input = user_input
            st.session_state.gpt_last_error = None
            st.session_state.gpt_chat_history.append({"role": "user", "content": user_input.strip()})
            try:
                with st.spinner("GPT думает..."):
                    response = openai.ChatCompletion.create(
                        model="gpt-3.5-turbo",
                        messages=st.session_state.gpt_chat_history,
                        temperature=0.7,
                        max_tokens=1024,
                    )
                    answer = response.choices[0].message.content.strip()
                    st.session_state.gpt_chat_history.append({"role": "assistant", "content": answer})
                    st.session_state.gpt_user_input = ""
            except Exception as e:
                st.session_state.gpt_last_error = f"Ошибка: {str(e)}"

        # --- Keep input in sync ---
        if not send_btn:
            st.session_state.gpt_user_input = user_input

        st.markdown("""
        <div style="height: 80px;"></div>
        """, unsafe_allow_html=True)

        st.info("Ваши сообщения не сохраняются и не отправляются третьим лицам. API-ключ хранится только в коде.")
    # ========== ВКЛАДКА 6: Инструкция ==========
    if tab == "📖 Инструкция":
        st.title("📖 Инструкция по использованию инструментов SEO-комбайна")
        st.markdown("""
        ### 🔎 Подбор ключей (Serpstat)
        **Автоматический сбор ключевых фраз для семантического ядра сайта.**
        - Загружает Excel-файл с названиями и URL.
        - Для каждой строки подбирает релевантные ключевые фразы через Serpstat API.
        - Результат можно скачать в виде таблицы.

        ---
        ### 🧙 Анализ уникальных слов
        **Находит слова, которые есть во втором тексте, но отсутствуют в первом (с учётом лемматизации).**
        - Вставьте два текста: первый — основной, второй — сравниваемый.
        - После анализа получите список уникальных слов с частотностью.

        ---
        ### 🔍 Подсветка слов + DOCX
        **Визуально выделяет ключевые слова в тексте и позволяет экспортировать результат в Word.**
        - Введите или загрузите текст, а также список ключевых слов (по одному на строку).
        - Ключевые слова будут подсвечены в тексте.
        - Можно скачать результат в формате DOCX.

        ---
        ### 🔍 SEO Meta Checker
        **Проверяет соответствие мета-тегов (Title/Description) и наличие ключевых фраз на страницах сайта.**
        - Загружает Excel-файл с URL, эталонными мета-тегами и фразами.
        - Сравнивает мета-теги сайта с эталонными, ищет точные и LSI-фразы в тексте страницы.
        - Доступна фильтрация, просмотр ошибок и скачивание результатов.

        ---
        ### 🧪 Tittle_Description +
        **Автоматическая сверка мета-тегов сайта с эталонными из таблицы с расчётом процента совпадения.**
        - Загружает Excel-файл с URL и мета-тегами.
        - Для каждого URL сравнивает мета-теги сайта с эталонными, показывает процент совпадения.
        - Можно фильтровать результаты и скачать итоговый файл.

        ---
        ### 🤖 GPT-ассистент
        **Чат-ассистент на базе OpenAI GPT.**
        - Позволяет вести диалог с искусственным интеллектом прямо в приложении.
        - Интерфейс в стиле мессенджера: история сообщений, ввод снизу, ответы сверху.
        - Можно очистить чат или скопировать последний ответ.

        ---
        **Если возникли вопросы — см. подсказки в каждой вкладке или обратитесь к разработчику.**
        """)
        st.info("Инструкция актуальна для всех вкладок приложения. Для подробностей см. описание внутри каждой вкладки.")
        return

    # ========== ВКЛАДКА 5: SEO Мета-Проверка 2.0 ==========
    if tab == "🧪 Tittle_Description +":
        import tempfile
        import altair as alt
        from difflib import SequenceMatcher

        st.set_page_config(page_title="SEO Мета-Проверка Apteka 9-1-1", layout="wide")
        st.title("🔍 SEO Мета-Проверка для сайта Apteka 9-1-1")

        st.markdown("""
        <style>
            .result-box {
                border: 1px solid #ccc;
                border-radius: 10px;
                padding: 15px;
                margin-bottom: 10px;
                background-color: #f9f9f9;
            }
            .highlight {
                font-weight: bold;
                color: #2c3e50;
            }
            .progress-bar {
                height: 20px;
                background-color: #f3f3f3;
                border-radius: 10px;
                margin: 10px 0;
            }
            .progress {
                height: 100%;
                background-color: #66bb6a;
                border-radius: 10px;
                transition: width 0.3s ease;
            }
            .tooltip {
                position: relative;
                display: inline-block;
                cursor: pointer;
            }
            .tooltip .tooltiptext {
                visibility: hidden;
                width: 200px;
                background-color: #555;
                color: #fff;
                text-align: center;
                padding: 5px;
                border-radius: 6px;
                position: absolute;
                z-index: 1;
                bottom: 125%;
                left: 50%;
                margin-left: -100px;
                opacity: 0;
                transition: opacity 0.3s;
            }
            .tooltip:hover .tooltiptext {
                visibility: visible;
                opacity: 1;
            }
            .stat-card {
                border: 1px solid #ddd;
                border-radius: 8px;
                padding: 15px;
                margin: 10px;
                background: white;
                box-shadow: 0 2px 4px rgba(0,0,0,0.1);
            }
        </style>
        """, unsafe_allow_html=True)

        def clean_text(text):
            text = str(text).lower()
            text = re.sub(r'%min_price%', '', text)
            text = re.sub(r'цена от [^|\n\r-]+', '', text)
            text = re.sub(r'цiна вiд [^|\n\r-]+', '', text)
            text = re.sub(r'[\-\|:,⭐⏩⚡🔹📦→®]', '', text)
            text = re.sub(r'грн|uah', '', text)
            text = re.sub(r'\s+', ' ', text)
            return text.strip()

        def get_similarity(text1, text2):
            return round(SequenceMatcher(None, text1, text2).ratio() * 100, 1)

        def get_meta_from_url(url):
            try:
                headers = {'User-Agent': 'Mozilla/5.0'}
                response = requests.get(url, headers=headers, timeout=10)
                soup = BeautifulSoup(response.text, 'html.parser')
                title = soup.title.string.strip() if soup.title else ''
                description = ''
                tag = soup.find("meta", attrs={"name": "description"})
                if tag and tag.get("content"):
                    description = tag["content"].strip()
                return title, description
            except:
                return '', ''

        tabs2 = st.tabs(["Загрузка", "Результаты"])

        with tabs2[0]:
            st.subheader("📄 Загрузка данных")
            st.markdown("""
            <div class="tooltip">
                Загрузите Excel-файл с данными
                <span class="tooltiptext">
                    Файл должен содержать колонки: URL, Title RU, Description RU, Title UA, Description UA
                </span>
            </div>
            """, unsafe_allow_html=True)
            uploaded_file = st.file_uploader("Загрузите Excel-файл с данными:", type=["xlsx"], key="meta2_file")
            if 'meta2_result_df' not in st.session_state:
                st.session_state.meta2_result_df = None

            if uploaded_file:
                df = pd.read_excel(uploaded_file)
                st.success("Файл загружен успешно. Начинаю проверку...")
                progress_bar = st.progress(0)
                status_text = st.empty()
                results = []
                total = len(df)
                for index, row in df.iterrows():
                    progress = (index + 1) / total
                    progress_bar.progress(progress)
                    status_text.text(f"Проверено {index + 1}/{total} URL")
                    url = str(row['URL']).strip()
                    url_ua = url.replace("apteka911.ua/", "apteka911.ua/ua/")
                    row_result = {
                        'URL': url,
                        'Title RU (таблица)': row.get('Title RU', ''),
                        'Description RU (таблица)': row.get('Description RU', ''),
                        'Title UA (таблица)': row.get('Title UA', ''),
                        'Description UA (таблица)': row.get('Description UA', '')
                    }
                    # RU
                    title_ru_site, desc_ru_site = get_meta_from_url(url)
                    row_result['Title RU (сайт)'] = title_ru_site
                    row_result['Description RU (сайт)'] = desc_ru_site
                    # UA
                    title_ua_site, desc_ua_site = get_meta_from_url(url_ua)
                    row_result['Title UA (сайт)'] = title_ua_site
                    row_result['Description UA (сайт)'] = desc_ua_site
                    # Сходство
                    title_ru_sim = get_similarity(clean_text(row.get('Title RU', '')), clean_text(title_ru_site))
                    desc_ru_sim = get_similarity(clean_text(row.get('Description RU', '')), clean_text(desc_ru_site))
                    title_ua_sim = get_similarity(clean_text(row.get('Title UA', '')), clean_text(title_ua_site))
                    desc_ua_sim = get_similarity(clean_text(row.get('Description UA', '')), clean_text(desc_ua_site))
                    row_result['Title RU Совпадение (%)'] = title_ru_sim
                    row_result['Description RU Совпадение (%)'] = desc_ru_sim
                    row_result['Title UA Совпадение (%)'] = title_ua_sim
                    row_result['Description UA Совпадение (%)'] = desc_ua_sim
                    results.append(row_result)
                status_text.text("Проверка завершена!")
                result_df = pd.DataFrame(results)
                st.session_state.meta2_result_df = result_df
                def highlight_percentage(p):
                    if p >= 80:
                        return f"✅ {p}%"
                    elif p >= 50:
                        return f"🟡 {p}%"
                    else:
                        return f"❌ {p}%"
                display_df = result_df.copy()
                for col in display_df.columns:
                    if "Совпадение" in col:
                        display_df[col] = display_df[col].apply(highlight_percentage)
                with tempfile.NamedTemporaryFile(delete=False, suffix=".xlsx") as tmp:
                    result_df.to_excel(tmp.name, index=False)
                    st.download_button("📥 Скачать результат", data=open(tmp.name, "rb"), file_name="seo_meta_check_result.xlsx")

        with tabs2[1]:
            st.subheader("📋 Результаты проверки")
            result_df = st.session_state.get('meta2_result_df', None)
            if result_df is not None:
                min_similarity = st.slider("Минимальный процент совпадения", 0, 100, 50)
                filtered_df = result_df.copy()
                filtered_df = filtered_df[(filtered_df['Title RU Совпадение (%)'] >= min_similarity) |
                                        (filtered_df['Description RU Совпадение (%)'] >= min_similarity) |
                                        (filtered_df['Title UA Совпадение (%)'] >= min_similarity) |
                                        (filtered_df['Description UA Совпадение (%)'] >= min_similarity)]
                def highlight_percentage(p):
                    if p >= 80:
                        return f"✅ {p}%"
                    elif p >= 50:
                        return f"🟡 {p}%"
                    else:
                        return f"❌ {p}%"
                display_df = filtered_df.copy()
                for col in display_df.columns:
                    if "Совпадение" in col:
                        display_df[col] = display_df[col].apply(highlight_percentage)
                st.dataframe(display_df, use_container_width=True)
            else:
                st.info("Загрузите файл для начала проверки")

    st.sidebar.markdown("---")
    st.sidebar.info("Сделано с заботой о твоём времени ❤️‍🔥", icon="💡")

    # ========== ВКЛАДКА 1: Подбор ключей (Serpstat) ==========
    if tab == "🔎 Подбор ключей (Serpstat)":
        st.title("🔎 Автоматизированный подбор ключей для семантического ядра")

        uploaded_file = st.file_uploader("Загрузи Excel-файл (.xlsx) с Названием, URL и Фразы", type=["xlsx"])
        start_button = st.button("🚀 Запустить обработку")

        if uploaded_file and start_button:
            with st.spinner("Обрабатываем файл..."):
                df = pd.read_excel(uploaded_file)
                if "Название" not in df.columns or "URL" not in df.columns or "Фразы в точном вхождении" not in df.columns:
                    st.error("В Excel должны быть столбцы 'Название', 'URL' и 'Фразы в точном вхождении'!")
                else:
                    df["Фразы в точном вхождении"] = ""
                    total = len(df)
                    progress_bar = st.progress(0)
                    status_text = st.empty()
                    for idx, row in df.iterrows():
                        name = str(row["Название"]).strip()
                        phrases = get_serpstat_phrases_top_filtered(name, top_n=10)
                        df.at[idx, "Фразы в точном вхождении"] = "\n".join(phrases) if phrases else "Нет данных"
                        progress = (idx + 1) / total
                        progress_bar.progress(progress)
                        status_text.text(f"Обработано {idx+1} из {total}...")
                    progress_bar.progress(1.0)
                    status_text.text("✅ Готово! Можно скачивать результат.")
                    st.success("Готово! Скачай Excel ниже 👇")
                    output = io.BytesIO()
                    df.to_excel(output, index=False)
                    st.download_button(
                        label="💾 Скачать результат",
                        data=output.getvalue(),
                        file_name="ключевые_фразы_по_таблице.xlsx",
                        mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
                    )

    # ========== ВКЛАДКА 2: Анализ уникальных слов ==========
    elif tab == "🧙 Анализ уникальных слов":
        st.title("🧙 Анализатор уникальных слов в тексте")
        st.info(
            "Этот инструмент находит слова, которые есть во втором тексте, но отсутствуют в первом. "
            "Анализ учитывает разные формы слов (падежи, числа), приводя их к начальной форме (лемме)."
        )

        if 'text1' not in st.session_state:
            st.session_state.text1 = ""
        if 'text2' not in st.session_state:
            st.session_state.text2 = ""

        def clear_text1():
            st.session_state.text1 = ""

        def clear_text2():
            st.session_state.text2 = ""

        col1, col2 = st.columns(2)
        with col1:
            st.header("Текст №1 (Основной)")
            st.session_state.text1 = st.text_area(
                label="Слова из этого текста будут исключены из анализа:",
                value=st.session_state.text1,
                height=300,
                placeholder="Вставьте сюда основной текст...",
                key="text_area1"
            )
            st.button("Очистить", on_click=clear_text1, key="clear_text1")

        with col2:
            st.header("Текст №2 (Сравниваемый)")
            st.session_state.text2 = st.text_area(
                label="Здесь будут искаться уникальные слова:",
                value=st.session_state.text2,
                height=300,
                placeholder="Вставьте сюда текст для сравнения...",
                key="text_area2"
            )
            st.button("Очистить", on_click=clear_text2, key="clear_text2")

        if st.button("🚀 Начать анализ", use_container_width=True):
            if st.session_state.text1 and st.session_state.text2:
                with st.spinner("Пожалуйста, подождите, идёт обработка текстов..."):
                    analysis_result = analyze_texts(st.session_state.text1, st.session_state.text2)
                st.header("Результат анализа")
                st.markdown("---")
                st.markdown(analysis_result)
            else:
                st.error("❗ Пожалуйста, введите тексты в оба поля для анализа.")

        st.markdown(
            """
            <style>
            .scroll-to-top {
                position: fixed;
                bottom: 20px;
                right: 20px;
                z-index: 1000;
            }
            </style>
            <button class="scroll-to-top" onclick="window.scrollTo({top: 0, behavior: 'smooth'})">↑ Наверх</button>
            """,
            unsafe_allow_html=True
        )

    # ========== ВКЛАДКА 3: Подсветка слов + DOCX ==========
    elif tab == "🔍 Подсветка слов + DOCX":
        st.title("🔍 Подсветка ключевых слов + Просмотр + Экспорт в Word")

        text_source = st.radio("Источник текста:", ["Ввести вручную", "Загрузить .txt файл"])

        if text_source == "Ввести вручную":
            input_text = st.text_area("✍️ Вставьте текст", height=300)
        else:
            uploaded_file = st.file_uploader("📄 Загрузите .txt файл", type=["txt"])
            input_text = uploaded_file.read().decode("utf-8") if uploaded_file else ""

        keyword_block = st.text_area("📋 Список слов (по одному на строку):", height=200)
        keywords = [w.strip().lower() for w in keyword_block.strip().splitlines() if w.strip()]

        if st.button("✨ Подсветить и показать результат"):
            if not input_text or not keywords:
                st.warning("Пожалуйста, введите текст и список слов.")
            else:
                html_text = input_text
                for word in keywords:
                    pattern = re.compile(rf'\b({re.escape(word)}\w*)\b', flags=re.IGNORECASE)
                    html_text = pattern.sub(r'<span style="background-color:yellow;">\1</span>', html_text)
                st.markdown("### 👀 Просмотр с подсветкой:")
                st.markdown(f"<div style='line-height:1.6'>{html_text}</div>", unsafe_allow_html=True)

                doc = Document()
                paragraph = doc.add_paragraph()
                words = re.split(r'(\W+)', input_text)
                for word in words:
                    clean_word = re.sub(r'\W+', '', word).lower()
                    match = any(re.fullmatch(rf'{re.escape(k)}\w*', clean_word, re.IGNORECASE) for k in keywords)
                    run = paragraph.add_run(word)
                    if match:
                        run.font.highlight_color = 7  # жёлтый
                buffer = BytesIO()
                doc.save(buffer)
                buffer.seek(0)
                b64 = base64.b64encode(buffer.read()).decode()
                href = f'<a href="data:application/octet-stream;base64,{b64}" download="highlighted.docx">📥 Скачать как DOCX</a>'
                st.markdown("### 💾 Скачай файл, если всё ок:")
                st.markdown(href, unsafe_allow_html=True)

    # ========== ВКЛАДКА 4: SEO Meta Checker ==========
    elif tab == "🔍 SEO Meta Checker":
        st.title("SEO Meta Checker для Apteka911")
        st.markdown("Проверка мета-тегов и фраз на веб-страницах сайта Apteka911.")
        st.markdown("""
        <style>
        .info-box { background-color: #f8f9fa; border: 1px solid #e0e0e0; border-radius: 5px; padding: 10px; margin: 10px 0; }
        .result-content { margin: 5px 0; padding: 8px; background: white; border: 1px solid #eee; border-radius: 3px; font-size: 0.95em; word-wrap: break-word; }
        .result-content b { font-weight: 600; }
        .stTabs [data-baseweb="tab-list"] { gap: 8px; }
        .stTabs [data-baseweb="tab-list"] .stTabs [data-baseweb="tab-list"] { gap: 4px; }
        .stTabs [data-baseweb="tab"] { height: 50px; padding: 0 25px; margin-right: 0; border-radius: 5px 5px 0 0; }
        .stTabs [aria-selected="true"] { background-color: #e6f2ff; border-bottom: 2px solid #007bff; }
        .stMetric { border: 1px solid #ddd; border-radius: 5px; padding: 10px; background-color: #f9f9f9;}
        </style>
        """, unsafe_allow_html=True)

        if 'debug_messages' not in st.session_state: st.session_state.debug_messages = []
        if RUSSIAN_STEMMER is None and 'stemmer_warning_shown' not in st.session_state:
            st.warning("Русский стеммер (PyStemmer) не инициализирован...")
            st.session_state.stemmer_warning_shown = True
        if not RAPIDFUZZ_AVAILABLE and 'rapidfuzz_warning_shown' not in st.session_state:
            st.warning("Библиотека rapidfuzz не найдена...")
            st.session_state.rapidfuzz_warning_shown = True

        if 'enable_lsi_truncation' not in st.session_state: st.session_state.enable_lsi_truncation = DEFAULT_ENABLE_LSI_TRUNCATION
        if 'lsi_trunc_max_remove' not in st.session_state: st.session_state.lsi_trunc_max_remove = DEFAULT_LSI_TRUNC_MAX_REMOVE
        if 'lsi_trunc_min_orig_len' not in st.session_state: st.session_state.lsi_trunc_min_orig_len = DEFAULT_LSI_TRUNC_MIN_ORIG_LEN
        if 'lsi_trunc_min_final_len' not in st.session_state: st.session_state.lsi_trunc_min_final_len = DEFAULT_LSI_TRUNC_MIN_FINAL_LEN
        if 'stem_fuzzy_ratio_threshold' not in st.session_state: st.session_state.stem_fuzzy_ratio_threshold = DEFAULT_STEM_FUZZY_RATIO_THRESHOLD

        debug_mode = st.sidebar.checkbox("🕵️ Включить режим отладки", value=st.session_state.get('debug_mode_main_cb', False), key="debug_mode_main_cb")

        if debug_mode:
            st.sidebar.markdown("--- Настройки LSI ---")
            if RAPIDFUZZ_AVAILABLE:
                st.session_state.stem_fuzzy_ratio_threshold = st.sidebar.slider("Порог схожести основ (%)",
                    min_value=50, max_value=100, value=st.session_state.stem_fuzzy_ratio_threshold, step=1, key="fuzzy_thresh_slider_ui")
            st.session_state.enable_lsi_truncation = st.sidebar.checkbox("Включить поиск по усечению LSI",
                value=st.session_state.enable_lsi_truncation, key="enable_trunc_cb")
            if st.session_state.enable_lsi_truncation:
                st.session_state.lsi_trunc_max_remove = st.sidebar.slider("Макс. удаляемых букв", 1, 5,
                    st.session_state.lsi_trunc_max_remove, key="trunc_max_rem_slider")
                st.session_state.lsi_trunc_min_orig_len = st.sidebar.slider("Мин. длина LSI слова для усечения", 5, 15,
                    st.session_state.lsi_trunc_min_orig_len, key="trunc_min_orig_slider")
                st.session_state.lsi_trunc_min_final_len = st.sidebar.slider("Мин. длина слова после усечения", 2, 5,
                    st.session_state.lsi_trunc_min_final_len, key="trunc_min_final_slider")

            st.sidebar.markdown("--- Отладка конкретного URL ---")
            st.session_state.manual_debug_url_val = st.sidebar.text_input("URL для сохранения HTML:",
                value=st.session_state.get("manual_debug_url_val", ""), key="manual_debug_url_input_field")
            st.session_state.manual_debug_lang_val = st.sidebar.selectbox("Язык для отладки URL:", ["ru", "ua"],
                index=['ru','ua'].index(st.session_state.get("manual_debug_lang_val", 'ru')), key="manual_debug_lang_select")
            if st.sidebar.button("Сохранить HTML для URL", key="save_html_btn") and st.session_state.manual_debug_url_val:
                st.sidebar.info(f"HTML для {st.session_state.manual_debug_lang_val.upper()}: {st.session_state.manual_debug_url_val}")
                get_page_data_for_lang(st.session_state.manual_debug_url_val, st.session_state.manual_debug_lang_val,
                                       debug_mode_internal=True,
                                       save_html_for_debug_manual=True, filename_prefix_manual="MANUAL_DEBUG_PAGE")
                st.sidebar.success("HTML (если получен) должен быть сохранен.")

        uploaded_file = st.file_uploader("📤 Загрузите Excel файл с данными", type=["xlsx"])
        if 'processed_data' not in st.session_state: st.session_state.processed_data = None

        if uploaded_file and st.button("🚀 Начать проверку всех URL из файла", key="start_full_processing_btn"):
            st.session_state.debug_messages = []
            df_excel = None
            try:
                df_excel = pd.read_excel(uploaded_file)
                st.subheader(f"Проверка содержимого файла: '{uploaded_file.name}'")
                st.markdown(f"Всего строк в файле: **{len(df_excel)}**. Первые 5 строк:")
                st.dataframe(df_excel.head())
                with st.expander("Поиск URL в загруженном файле (для проверки наличия)", expanded=False):
                    url_to_search_in_df = st.text_input("Введите часть URL для поиска в таблице:", key="df_url_search_input")
                    if url_to_search_in_df:
                        if COL_URL_RU_EXCEL not in df_excel.columns: st.error(f"Колонка '{COL_URL_RU_EXCEL}' не найдена в вашем Excel файле!")
                        else:
                            search_results_df = df_excel[df_excel[COL_URL_RU_EXCEL].astype(str).str.contains(url_to_search_in_df, case=False, na=False)]
                            if not search_results_df.empty: st.write(f"Найдены строки с '{url_to_search_in_df}':"); st.dataframe(search_results_df)
                            else: st.warning(f"URL, содержащий '{url_to_search_in_df}', не найден в загруженном файле.")
                if COL_URL_RU_EXCEL not in df_excel.columns:
                    st.error(f"В файле отсутствует ОБЯЗАТЕЛЬНАЯ колонка '{COL_URL_RU_EXCEL}'!")
                    if debug_mode: display_debug_messages()
                    return

                load_progress_bar_ui = st.progress(0.0)
                load_progress_text_ui = st.empty()
                st.session_state.global_progress_text = "Инициализация загрузки..."
                st.session_state.global_progress_value = 0.0
                load_progress_text_ui.info(st.session_state.global_progress_text)

                st.session_state.processed_data = load_all_pages_data_for_both_langs(df_excel, debug_mode)

                load_progress_text_ui.success(st.session_state.get('global_progress_text', "Загрузка данных завершена!"))
                load_progress_bar_ui.progress(st.session_state.get('global_progress_value', 1.0))

                urls_total_count = len(df_excel)
                urls_load_errors_count_ru = sum(1 for data_dict in st.session_state.processed_data.values() if data_dict.get('ru', {}).get('error'))
                urls_load_errors_count_ua = sum(1 for data_dict in st.session_state.processed_data.values() if data_dict.get('ua', {}).get('error'))
                st.subheader(f"📊 Общая сводка по загрузке страниц")
                summary_cols = st.columns(2)
                summary_cols[0].metric("Всего URL в файле для обработки", urls_total_count)
                summary_cols[1].metric("URL с ошибками загрузки (RU)", urls_load_errors_count_ru, delta_color="inverse" if urls_load_errors_count_ru > 0 else "off")
                summary_cols[1].metric("URL с ошибками загрузки (UA)", urls_load_errors_count_ua, delta_color="inverse" if urls_load_errors_count_ua > 0 else "off")

            except pd.errors.EmptyDataError: st.error("Ошибка: Excel файл пуст."); st.session_state.processed_data = None
            except KeyError as e: st.error(f"Ошибка: Отсутствует колонка '{str(e)}' в Excel."); st.session_state.processed_data = None
            except Exception as e:
                st.error(f"Критическая ошибка при загрузке или начальной обработке файла: {str(e)}")
                st.session_state.processed_data = None
                if 'debug_mode' in locals() and debug_mode:
                    st.exception(e)
                else:
                    print(f"Критическая ошибка (debug_mode не был определен или False в момент исключения): {e}")
                    traceback.print_exc()

        if uploaded_file and st.session_state.get('processed_data') is not None:
            df_for_tabs_display = None
            try:
                if 'df_excel' in locals() and df_excel is not None:
                    df_for_tabs_display = df_excel
                else:
                    df_for_tabs_display = pd.read_excel(uploaded_file)
            except Exception as e_read_tabs:
                st.error(f"Не удалось подготовить данные для отображения вкладок: {e_read_tabs}")
                df_for_tabs_display = None

            if df_for_tabs_display is not None:
                main_ru_tab, main_ua_tab = st.tabs(["🇷🇺 Русская Версия", "🇺🇦 Украинская Версия"])
                with main_ru_tab:
                    run_checks_for_language('ru', df_for_tabs_display, st.session_state.processed_data, debug_mode)
                with main_ua_tab:
                    run_checks_for_language('ua', df_for_tabs_display, st.session_state.processed_data, debug_mode)

        if 'debug_mode' in locals() and debug_mode:
            display_debug_messages()

    if tab == "🖼️ Проверка URL изображений аптек":
        pharmacy_image_url_checker_tab()

if __name__ == "__main__":
    main()